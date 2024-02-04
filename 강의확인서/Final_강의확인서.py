import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
import openpyxl
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
import os 


def browse_file(entry):
    file_path = filedialog.askopenfilename(filetypes=[("Excel 파일", "*.xlsm")])
    if file_path:  # 사용자가 파일을 선택한 경우
        entry.delete(0, tk.END)
        entry.insert(0, file_path)

def browse_folder(entry):
    folder_selected = filedialog.askdirectory(initialdir="C:/")
    if folder_selected:  # 사용자가 폴더를 선택한 경우
        entry.delete(0, tk.END)
        entry.insert(0, folder_selected)

def show_message(message):
    messagebox.showinfo("Message", message)

def start_clicked():
    # "시작" 버튼 클릭 시, 각 위젯에서 값을 읽어옴
    global MONTH, YEAR, DAY, LECTURE_NAME, LECTURE_SHEETNAME, TOPIC, RAW_DATA, SAVE_FOLDER, TYPE

    MONTH = int(month_combobox.get())
    YEAR = year_spinbox.get()
    DAY = day_combobox.get()
    LECTURE_NAME = lecture_name_entry.get()
    LECTURE_SHEETNAME = lecture_sheetname_entry.get()
    TOPIC = topic_entry.get()
    RAW_DATA = raw_data_entry.get()
    SAVE_FOLDER = save_folder_entry.get()
    TYPE = type_combobox.get()


    if not all([MONTH, YEAR, DAY, LECTURE_NAME, LECTURE_SHEETNAME, TOPIC, RAW_DATA, SAVE_FOLDER, TYPE]):
        show_message("입력값을 확인하세요.")
        return 
    # 주어진 코드 실행

    generate_confirmation_doc()

 
def generate_confirmation_doc():
    # 이 부분에 문서 생성 코드 추가
    print(f'YEAR: {YEAR}, MONTH: {MONTH}, DAY: {DAY} {RAW_DATA}, {SAVE_FOLDER}')
    print(f'LECTURE_NAME: {LECTURE_NAME}, LECTURE_SHEETNAME: {LECTURE_SHEETNAME}, TOPIC: {TOPIC}')

    save_file = f'{MONTH}월_{LECTURE_SHEETNAME}'
    if TYPE == "해 설 사":
        save_file += '해설확인서.docx'
        title_text = '해  설  확  인  서'
    elif TYPE == "강  사":
        save_file += '강의확인서.docx'
        title_text = '강  의  확  인  서'
    else:
        show_message("잘못된 TYPE 값입니다. 해설과 강의 중에 선택해주세요")
        return

    save_data = os.path.join(SAVE_FOLDER, save_file)

    wb = openpyxl.load_workbook(RAW_DATA, read_only=True)
    lec_raw_sheet = wb['통합시트']

    first_row_number = 2
    last_row_number = lec_raw_sheet.max_row

    PROGRAM_DATA = {}

    for row in lec_raw_sheet.iter_rows(min_row=first_row_number, max_row=last_row_number):
        date = row[1].value


        if date is not None:
            if date.month != MONTH:
                continue

            lec_name = row[0].value

            date = row[1].value.strftime("%Y-%m-%d")
            dayofweek = row[2].value
            timetable = row[3].value  # 시작시각:시작분 ~ 종료시각:종료분
            applicant = row[4].value
            division = row[5].value
            teacher1 = row[8].value
            teacher2 = row[9].value

            # timetable이 예상대로 구성되어 있는지 확인
            if "~" not in timetable:
                messagebox.showwarning("경고", f"{date}의 시간에 ~가 없습니다.{timetable}")
                if root.winfo_exists():
                    root.destroy()  # GUI 닫기
                return  # Exit the function



            # 시작시각과 종료시각을 추출하여 각각의 시간 데이터로 변환
            start_time_str, end_time_str = timetable.split("~")
            start_time = datetime.strptime(start_time_str.strip(), "%H:%M")
            end_time = datetime.strptime(end_time_str.strip(), "%H:%M")



            # 데이터에 추가
            data = [date, dayofweek, start_time, end_time, applicant, division]


            if not lec_name in PROGRAM_DATA:
                PROGRAM_DATA[lec_name] = {}

            for teacher in (teacher1, teacher2):
                if teacher is None:
                    continue
                if teacher == '-':
                    continue

                if not teacher in PROGRAM_DATA[lec_name]:
                    PROGRAM_DATA[lec_name][teacher] = []

                PROGRAM_DATA[lec_name][teacher].append(data)


    ws_personal = wb["인적사항"]
    PERSONAL_DATA = {}

    for row in ws_personal.iter_rows(min_row=2):
        name = row[0].value

        if not name:
            break

        position = row[1].value
        social_id = row[2].value
        address = row[3].value
        phone = row[4].value
        bank = row[5].value
        banknum = row[6].value

        PERSONAL_DATA[name] = [position, social_id, address, phone, bank, banknum]

    wb.close()

    

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')

    first_teacher = True

    for teacher_name, teacher_data in PROGRAM_DATA[LECTURE_SHEETNAME].items():
        if teacher_name not in PERSONAL_DATA:
            # 해당 강사의 인적사항이 없을 때 메시지 박스를 띄우고 GUI를 닫음
            messagebox.showwarning("경고", f"해당 강사({teacher_name})의 인적사항이 없습니다.")
            if root.winfo_exists():
                root.destroy()  # GUI 닫기
            return  # Exit the function


        if first_teacher:
            first_teacher = False
        else:
            doc.add_page_break()
    

        # "강 의 확 인 서"를 중앙 정렬하고 글씨 크기를 15로 변경
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_paragraph.runs[0]
        run.font.size = Pt(20)
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        

   
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER  # 빈 줄 추가
        doc.add_paragraph(f'1. 사 업 명: {LECTURE_NAME}')
        doc.add_paragraph('2. 강의일시 및 대상')

        # 시작시각을 기준으로 정렬
        sorted_teacher_data = sorted(
            teacher_data,
            key=lambda x: (x[0], x[2])
        )

        # 표 스타일 적용

        table = doc.add_table(rows=1, cols=5)
        table.style = doc.styles['Table Grid']
        table.autofit = False

        # 첫 번째 열에 대한 헤더 추가
        cell = table.cell(0, 0)
        cell.text = "순번"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for col_num, header_text in enumerate(['날짜(요일)', '강의시간', '신청기관', '반']):
            cell = table.cell(0, col_num + 1)
            cell.text = header_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            # 순차 번호를 초기화
            sequence_number = 1

        for i, data in enumerate(sorted_teacher_data, start=1):
            row_cells = table.add_row().cells

            # 첫 번째 열(순번)에 값을 설정
            row_cells[0].text = str(sequence_number)
            sequence_number += 1

            # 다른 열에 값을 설정
            row_cells[1].text = f'{data[0]}({data[1]})'
            row_cells[2].text = f'{data[2].strftime("%H:%M")} ~ {data[3].strftime("%H:%M")}'
            row_cells[3].text = data[4]
            row_cells[4].text = data[5]
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '나눔고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕') 

        
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER  # 빈 줄 추가


    
        doc.add_paragraph(f'3. 강의주제 : {TOPIC}')
        doc.add_paragraph(f'4. 강    사 : {teacher_name}')    
        doc.add_paragraph(f'    ○ 소속및직위 : {PERSONAL_DATA[teacher_name][0]}')
        doc.add_paragraph(f'    ○ 생년월일 : {PERSONAL_DATA[teacher_name][1]}')
        doc.add_paragraph(f'    ○ 주    소 : {PERSONAL_DATA[teacher_name][2]}')
        doc.add_paragraph(f'    ○ 연 락 처 : {PERSONAL_DATA[teacher_name][3]}')
        doc.add_paragraph(f'    ○ 계좌번호 : {PERSONAL_DATA[teacher_name][4]} {PERSONAL_DATA[teacher_name][5]}')

        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER  # 빈 줄 추가
        doc.add_paragraph(f'{YEAR}. {MONTH}. {DAY}. ').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER  # 빈 줄 추가

        blank_name = '  '.join(teacher_name)
        doc.add_paragraph(f'{TYPE}    {blank_name}   (서 명)').alignment = WD_ALIGN_PARAGRAPH.CENTER


        # 표 스타일 적용
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕') 
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER                  
    # 문서 저장
    doc.save(save_data)

    try:
        # Save the document
        doc.save(save_data)
        messagebox.showinfo("알림", "강의확인서가 성공적으로 생성되었습니다.")
        return True  # Document saved successfully
    except Exception as e:
        # If an exception occurs during document saving, show a warning
        messagebox.showwarning("경고", f"강의확인서 저장 중 오류가 발생했습니다: {str(e)}")
        return False  # Document not saved successfully



# GUI 생성
root = tk.Tk()
root.title("강의확인서 출력")

# 각각의 값 입력 위젯 생성
# 연도는 스핀 박스로 변경, 기본값은 2023

type_frame = LabelFrame(root, text="해설/강의")
type_label = ttk.Label(type_frame, text="타입:")
type_combobox = ttk.Combobox(type_frame, values=["강  사", "해 설 사"], width=10)

year_frame = LabelFrame(root, text="날짜지정")
year_label = ttk.Label(year_frame, text="연도:", width=5)
year_spinbox = ttk.Spinbox(year_frame, from_=2023, to=2050, width=10)

month_label = ttk.Label(year_frame, text="교육월:", width=5)
month_combobox = ttk.Combobox(year_frame, values=[f'{i}' for i in range(1, 13)], width=5)

day_label = ttk.Label(year_frame, text="발행일:", width=5)
day_combobox = ttk.Combobox(year_frame, values=[str(i) for i in range(1, 32)], width=5)

lecture_name_label = ttk.Label(root, text="강의확인서 사업명(예:섬강따라물빛여행):")
lecture_name_entry = ttk.Entry(root, width=35)

lecture_sheetname_label = ttk.Label(root, text="엑셀 강의명(예: 섬강):")
lecture_sheetname_entry = ttk.Entry(root, width=35)

topic_label = ttk.Label(root, text="강의주제(물절약):")
topic_entry = ttk.Entry(root, width=35)

raw_data_entry = ttk.Entry(root, width=35)
raw_data_button = ttk.Button(root, text="교육파일열기", width=15, command=lambda: browse_file(raw_data_entry))

save_folder_entry = ttk.Entry(root,width=35)
save_folder_button = ttk.Button(root, text="저장경로설정", width=15, command=lambda: browse_folder(save_folder_entry))

# "시작" 버튼 생성
start_button = ttk.Button(root, text="시작", command=start_clicked)

# 각 위젯 배치;
type_frame.grid(column=0, row=0, sticky=tk.W, padx=1, pady=1)
type_label.grid(column=0, row=0, sticky=tk.W, padx=5, pady=5)
type_combobox.grid(column=1, row=0, padx=5, pady=5)


#year_frame에 비치
year_frame.grid(column=0, row=1, sticky= tk.W, padx=1, pady=1)
year_label.grid(column=0, row=1, sticky= tk.W, padx=5, pady=5)
year_spinbox.grid(column=1, row=1, padx=5, pady=5)

month_label.grid(column=0, row=2, sticky= tk.W, padx=5, pady=5)
month_combobox.grid(column=1, row=2, padx=5, pady=5)

day_label.grid(column=0, row=3, sticky= tk.W, padx=5, pady=5)
day_combobox.grid(column=1, row=3, padx=5, pady=5)

lecture_name_label.grid(column=0, row=6, sticky=tk.W, padx=5, pady=5)
lecture_name_entry.grid(column=1, row=6, padx=5, pady=5)

lecture_sheetname_label.grid(column=0, row=7, sticky=tk.W, padx=5, pady=5)
lecture_sheetname_entry.grid(column=1, row=7, padx=5, pady=5)

topic_label.grid(column=0, row=8, sticky=tk.W, padx=5, pady=5)
topic_entry.grid(column=1, row=8, padx=5, pady=5)

raw_data_entry.grid(column=1, row=4,  padx=1, pady=1)
raw_data_button.grid(column=0, row=4,  padx=5, pady=5, sticky=tk.W)

save_folder_entry.grid(column=1, row=5, padx=3, pady=5)
save_folder_button.grid(column=0, row=5, padx=5, pady=5,sticky=tk.W)

start_button.grid(column=0, row=9, columnspan=2, pady=10)

root.mainloop()
